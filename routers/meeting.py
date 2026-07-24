"""
routers/meeting.py — ELN Meeting Copilot
POST /api/ai/speech              — audio transcription via Azure AI Speech
                                    Fast Transcription REST API (external
                                    contract unchanged: {"transcript":...,
                                    "language_detected":...} or pass-through
                                    of a text transcript)
POST /api/ai/meeting              — transcript -> structured report -> Word
                                     doc -> blob. Synchronous. External
                                     contract unchanged.
POST /api/ai/meeting/chunk        — transcribe ONE short (~4 min) segment of
                                     a live in-progress recording. Fast,
                                     durable, synchronous per-chunk call —
                                     this is what lets long meetings survive
                                     without one giant inline transcription.
POST /api/ai/meeting/upload       — one whole pre-recorded file. Returns
                                     immediately with {report_id, status:
                                     "processing"}; transcription + report
                                     generation run as a BackgroundTask.
GET  /api/ai/meeting/status/{id}  — poll target for /api/ai/meeting/upload.
                                     Once status != "processing", also
                                     returns the sidecar JSON fields
                                     (candidate_notes, email_draft).
GET  /api/ai/meeting/reports      — list all meeting reports

═══════════════════════════════════════════════════════════════════════════
SPEECH TRANSCRIPTION — CONFIRMED LIVE, 2026-07-21, against speech-eln-covvalent
═══════════════════════════════════════════════════════════════════════════
The previous implementation used the Azure Speech SDK's synchronous
recognize_once() (see git history) — replaced entirely with a plain httpx
call to the Fast Transcription REST API. Verified via a real Cloud Shell
round-trip against the live resource, not assumed:

- ENDPOINT: the resource has NO custom subdomain
  (customSubDomainName is null — `az cognitiveservices account show`
  confirmed this), so the "<name>.cognitiveservices.azure.com" hostname
  guessed initially does not exist in DNS at all (curl: "Could not resolve
  host"). The real, working base is the REGIONAL endpoint:
  https://centralindia.api.cognitive.microsoft.com — confirmed via
  `properties.endpoint` on the resource itself, matching SPEECH_REGION
  below. Auth is via Ocp-Apim-Subscription-Key (the existing
  _get_speech_key() managed-identity -> ARM key-retrieval flow, unchanged),
  NOT an Entra bearer token — so the missing custom subdomain doesn't
  actually block anything here (that requirement is specific to Entra-token
  auth against Cognitive Services, not subscription-key auth).
- API-VERSION: 2024-11-15 confirmed working (HTTP 422 InvalidAudioFormat on
  a synthetic tone, then a full HTTP 200 with real phrases on a genuine
  speech clip). 2025-10-15 behaved identically wherever tested; 2024-11-15
  was chosen as the stable, documented GA baseline. 2024-05-15-preview was
  tested and explicitly REJECTED for this use case — it 400s with "Exactly
  one input locale must be specified, single language id is not yet
  supported," i.e. that preview version does not support the multi-locale
  array Hinglish detection depends on.
- REQUEST SHAPE: multipart/form-data, "audio" file field + "definition"
  field containing a JSON string `{"locales": [...]}` — confirmed accepted
  with locales=["en-IN","hi-IN"] (a 2-locale array), i.e. multi-locale
  auto-detection is supported on 2024-11-15.
- RESPONSE SHAPE: confirmed via a real Hinglish test clip —
  {"durationMilliseconds": int,
   "combinedPhrases": [{"text": "<whole transcript>"}],
   "phrases": [{"offsetMilliseconds": int, "durationMilliseconds": int,
                "text": "...", "words": [{"text":...,
                "offsetMilliseconds":...,"durationMilliseconds":...}],
                "locale": "en-IN", "confidence": 0.81...}, ...]}
  The per-phrase "locale" field is real and populated — confirmed the key
  name is literally "locale", not "language" or anything else guessed.
  (Observed behavior, not a bug to fix: on the live Hinglish test clip,
  every phrase came back tagged "en-IN" even though much of the content
  was Hindi words in Latin script — the recognizer leaned entirely to
  en-IN rather than alternating per phrase. The per-phrase "locale" field
  is read and passed through regardless of what value the model actually
  returns in practice.)

BLOB STORAGE — reuses the exact "auto-create container if missing, app
identity already has Storage Blob Data Contributor at account scope"
pattern already used for eln-reports / eln-chat-uploads elsewhere in this
codebase (see _upload_blob() below, now parameterized by container so it's
shared across eln-reports and the new eln-meeting-audio container — no new
role assignment requested or required).

AUTHOR/USER SCOPING — every blob key written by this file (report docx,
JSON sidecar, chunk audio, whole-file audio archive) is namespaced by a
slugified author identifier, so concurrent meetings from different
@covvalent.com users never collide. Job status is tracked entirely via the
`status` column on the existing eln_meeting_reports table, keyed by
report_id (a DB identity column) — there is deliberately NO in-memory/
global job-state dict anywhere in this file: report_id already uniquely
scopes each job, and a DB-backed status additionally survives across
worker processes/restarts, which an in-memory dict would not.

SIDECAR JSON — per instructions, no migration and no new SQL columns.
candidate_notes / email_draft / per-chunk phrase-locale metadata are written
to a JSON file next to the report docx blob: same container, same key
prefix, ".json" suffix instead of ".docx". See _upload_blob() call sites.
"""

import os
import io
import re
import json
import time
import uuid
import asyncio
import logging
import httpx
import pyodbc
from datetime import datetime, date
from typing import Optional
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks
from pydantic import BaseModel

router = APIRouter()
logger = logging.getLogger(__name__)

# ── Config ────────────────────────────────────────────────────────────────────
AOAI_ENDPOINT    = os.environ.get("AOAI_ENDPOINT", "https://aoai-eln-covvalent-2e2ec.openai.azure.com")
AOAI_DEPLOYMENT  = "gpt-4o"
AOAI_API_VERSION = "2024-12-01-preview"

# Speech — see the module docstring's "CONFIRMED LIVE" section for how each
# of these three values was verified against the real speech-eln-covvalent
# resource (not assumed).
SPEECH_REGION      = os.environ.get("SPEECH_REGION", "centralindia")
SPEECH_RESOURCE    = os.environ.get("SPEECH_RESOURCE", "speech-eln-covvalent")
SPEECH_ENDPOINT    = os.environ.get("SPEECH_ENDPOINT", f"https://{SPEECH_REGION}.api.cognitive.microsoft.com")
FAST_TRANSCRIBE_API_VERSION = "2024-11-15"
MEETING_LOCALES    = ["en-IN", "hi-IN"]  # Hinglish: per-phrase locale auto-detection, not one language for the whole file

SUBSCRIPTION_ID  = "9e25d11c-3753-4b8c-a575-0bcc44f964d4"
RESOURCE_GROUP   = "rg-eln-covvalent"
STORAGE_ACCOUNT  = "stelncoovalent"
BLOB_CONTAINER   = "eln-reports"
AUDIO_CONTAINER  = "eln-meeting-audio"

_EXT_MAP = {
    "audio/wav":     ".wav",
    "audio/webm":    ".webm",
    "video/webm":    ".webm",
    "audio/mp4":     ".mp4",
    "audio/x-m4a":   ".mp4",
    "audio/aac":     ".aac",
    "audio/ogg":     ".ogg",
    "audio/mpeg":    ".mp3",
}

# ── Pydantic models ───────────────────────────────────────────────────────────
class MeetingRequest(BaseModel):
    transcript:   str
    project_code: Optional[str] = None
    author:       str = ""

class ChunkResponse(BaseModel):
    session_id:  str
    chunk_index: int
    text:        str
    locale:      Optional[str] = None

class MeetingUploadResponse(BaseModel):
    report_id: int
    status:    str

# ── DB helpers ────────────────────────────────────────────────────────────────
def _get_conn():
    conn_str = (
        f"DRIVER={{ODBC Driver 18 for SQL Server}};"
        f"SERVER={os.environ['SQL_SERVER']};"
        f"DATABASE={os.environ['SQL_DATABASE']};"
        f"UID={os.environ['SQL_USERNAME']};"
        f"PWD={os.environ['SQL_PASSWORD']};"
        f"TrustServerCertificate=yes;"
        f"Connection Timeout=30;"
    )
    return pyodbc.connect(conn_str)

def _serialize(v):
    if isinstance(v, (datetime, date)):
        return v.isoformat()
    return v

def _clean(rows):
    return [{k: _serialize(v) for k, v in row.items()} for row in rows]

def _rows_to_dicts(cur):
    cols = [c[0] for c in cur.description]
    return [dict(zip(cols, row)) for row in cur.fetchall()]

# ── Project directory (product-name → project_code resolution) ──────────────
def _get_project_directory(conn) -> list:
    """Returns [{project_code, product_name, generic_name, cas_number}, ...]
    for every project, for product-name-to-project_code resolution."""
    cur = conn.cursor()
    cur.execute("""
        SELECT project_code, title AS product_name, generic_name, cas_number
        FROM eln_projects
        ORDER BY project_code
    """)
    return _clean(_rows_to_dicts(cur))

_directory_cache = {"data": None, "fetched_at": 0}
_DIRECTORY_TTL_SECONDS = 600  # 10 minutes

def get_project_directory_cached(conn) -> list:
    now = time.time()
    if _directory_cache["data"] is None or (now - _directory_cache["fetched_at"]) > _DIRECTORY_TTL_SECONDS:
        _directory_cache["data"] = _get_project_directory(conn)
        _directory_cache["fetched_at"] = now
    return _directory_cache["data"]

# ── Speech key via managed identity ──────────────────────────────────────────
def _get_speech_key() -> str:
    """Retrieve Cognitive Services key using managed identity credentials."""
    from azure.identity import ManagedIdentityCredential
    from azure.mgmt.cognitiveservices import CognitiveServicesManagementClient
    cred   = ManagedIdentityCredential()
    client = CognitiveServicesManagementClient(cred, SUBSCRIPTION_ID)
    keys   = client.accounts.list_keys(RESOURCE_GROUP, SPEECH_RESOURCE)
    return keys.key1

# Cached with a TTL — chunk transcription now calls this far more often
# (every ~4 min per active recording) than the old one-shot-per-meeting SDK
# flow did, so avoid re-hitting the ARM key-list API on every single chunk.
_speech_key_cache = {"key": None, "fetched_at": 0}
_SPEECH_KEY_TTL_SECONDS = 1800  # 30 minutes

def _get_speech_key_cached() -> str:
    now = time.time()
    if _speech_key_cache["key"] is None or (now - _speech_key_cache["fetched_at"]) > _SPEECH_KEY_TTL_SECONDS:
        _speech_key_cache["key"] = _get_speech_key()
        _speech_key_cache["fetched_at"] = now
    return _speech_key_cache["key"]

# ── Fast Transcription REST call ─────────────────────────────────────────────
async def _transcribe_via_rest(audio_bytes: bytes, filename: str, content_type: str = None, timeout: float = 180.0) -> dict:
    """
    POST to the Azure AI Speech Fast Transcription REST API — see the module
    docstring's "CONFIRMED LIVE" section for exactly how endpoint,
    api-version, and request/response shape were verified against the real
    speech-eln-covvalent resource.
    """
    loop = asyncio.get_event_loop()
    speech_key = await loop.run_in_executor(None, _get_speech_key_cached)

    files = {"audio": (filename, audio_bytes, content_type or "application/octet-stream")}
    data  = {"definition": json.dumps({"locales": MEETING_LOCALES})}

    async with httpx.AsyncClient(timeout=timeout) as client:
        r = await client.post(
            f"{SPEECH_ENDPOINT}/speechtotext/transcriptions:transcribe",
            params={"api-version": FAST_TRANSCRIBE_API_VERSION},
            headers={"Ocp-Apim-Subscription-Key": speech_key},
            files=files,
            data=data,
        )
        r.raise_for_status()
        return r.json()


def _extract_transcript_and_locales(result: dict) -> tuple:
    """
    Returns (full_text, phrase_locales) from a confirmed-shape Fast
    Transcription response — see module docstring. full_text comes from
    combinedPhrases (falls back to joining phrases[].text if
    combinedPhrases is absent/empty). phrase_locales is the real per-phrase
    locale detail (field name "locale", confirmed live), not one language
    guessed for the whole recording.
    """
    combined  = result.get("combinedPhrases") or []
    full_text = " ".join(p.get("text", "") for p in combined if p.get("text")).strip()

    phrases = result.get("phrases") or []
    if not full_text:
        full_text = " ".join(p.get("text", "") for p in phrases if p.get("text")).strip()

    phrase_locales = [
        {
            "text":               p.get("text", ""),
            "locale":             p.get("locale"),
            "offsetMilliseconds": p.get("offsetMilliseconds"),
        }
        for p in phrases
    ]
    return full_text, phrase_locales


async def _transcribe_audio(audio_bytes: bytes, content_type: str) -> tuple:
    """
    Transcribe audio via the Fast Transcription REST API. Returns
    (text, language_detected) — matches /api/ai/speech's existing external
    contract. language_detected is the first non-null per-phrase locale
    found (a real detected value now, e.g. "en-IN"), falling back to
    "unknown" if no phrases came back (e.g. silence).
    """
    ct  = (content_type or "").split(";")[0].strip().lower()
    ext = _EXT_MAP.get(ct, ".wav")

    result = await _transcribe_via_rest(audio_bytes, f"audio{ext}", content_type=ct)
    full_text, phrase_locales = _extract_transcript_and_locales(result)
    lang = next((p["locale"] for p in phrase_locales if p.get("locale")), "unknown")
    return full_text, lang

# ── GPT call ──────────────────────────────────────────────────────────────────
async def _call_gpt(system_prompt: str, user_content: str) -> str:
    from azure.identity.aio import DefaultAzureCredential
    cred  = DefaultAzureCredential()
    token = await cred.get_token("https://cognitiveservices.azure.com/.default")
    await cred.close()

    async with httpx.AsyncClient(timeout=120.0) as client:
        r = await client.post(
            f"{AOAI_ENDPOINT}/openai/deployments/{AOAI_DEPLOYMENT}/chat/completions"
            f"?api-version={AOAI_API_VERSION}",
            headers={"Authorization": f"Bearer {token.token}", "Content-Type": "application/json"},
            json={
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": user_content},
                ],
                "max_tokens":  4096,
                "temperature": 0.3,
            },
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]

# ── Parse GPT section output ──────────────────────────────────────────────────
SECTION_RE = re.compile(
    r"##\s*(CONTEXT|HYPOTHESIS REVIEW|EXPERIMENT PLAN|RISKS?|CANDIDATE_NOTES|EMAIL_DRAFT)\s*\n",
    re.IGNORECASE
)

_SECTION_KEY_MAP = {
    "context":            "context",
    "hypothesis review":  "hypotheses",
    "experiment plan":    "experiment_plan",
    "risk":               "risks",
    "risks":              "risks",
    "candidate_notes":    "candidate_notes",
    "email_draft":        "email_draft",
}

def _parse_sections(gpt_text: str) -> dict:
    parts = SECTION_RE.split(gpt_text)
    sections: dict = {}
    if len(parts) > 2:
        for i in range(1, len(parts) - 1, 2):
            header  = parts[i].lower().strip().rstrip("s")
            content = parts[i + 1].strip()
            key = _SECTION_KEY_MAP.get(parts[i].lower().strip(), parts[i].lower().strip())
            sections[key] = content
    if not sections:
        sections["full"] = gpt_text
    return sections

def _extract_topic(gpt_text: str) -> str:
    m = re.search(r"##\s*CONTEXT\s*\n+(.+)", gpt_text, re.IGNORECASE)
    if m:
        return m.group(1).strip()[:60]
    lines = [l.strip() for l in gpt_text.split("\n") if l.strip() and not l.startswith("#")]
    return lines[0][:60] if lines else "meeting"

def _slug(text: str, max_len: int = 40) -> str:
    clean = re.sub(r"[^\w\s-]", "", text).strip()
    return re.sub(r"\s+", "-", clean).lower()[:max_len]

def _strip_json_fence(text: str) -> str:
    """Strip a ```json ... ``` (or bare ``` ... ```) fence around a GPT section."""
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"```\s*$", "", text)
    return text.strip()

def _parse_candidate_notes_and_email(sections: dict) -> tuple:
    """Shared by the synchronous /api/ai/meeting endpoint and the
    /api/ai/meeting/upload background task — same parsing, one place."""
    candidate_notes = []
    raw_candidates  = sections.get("candidate_notes", "")
    if raw_candidates:
        try:
            parsed = json.loads(_strip_json_fence(raw_candidates))
            if isinstance(parsed, list):
                candidate_notes = [item for item in parsed if isinstance(item, dict)]
        except (json.JSONDecodeError, ValueError) as e:
            print(f"[MEETING WARNING] Failed to parse CANDIDATE_NOTES JSON: {e}")

    email_draft = {"subject": "", "body": ""}
    raw_email   = sections.get("email_draft", "")
    if raw_email:
        try:
            parsed = json.loads(_strip_json_fence(raw_email))
            if isinstance(parsed, dict):
                email_draft = {
                    "subject": parsed.get("subject", ""),
                    "body":    parsed.get("body", ""),
                }
        except (json.JSONDecodeError, ValueError) as e:
            print(f"[MEETING WARNING] Failed to parse EMAIL_DRAFT JSON: {e}")

    return candidate_notes, email_draft

# ── Word doc builder ──────────────────────────────────────────────────────────
def _set_cell_shading(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _build_meeting_report_docx(
    topic:          str,
    project_code:   Optional[str],
    author:         str,
    gpt_sections:   dict,
    generated_date: str,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    RGB = RGBColor

    doc    = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # Footer
    fp = doc.sections[0].footer.paragraphs[0]
    fp.text      = (
        f"Covvalent / Rainboweucalyptus Technologies Pvt. Ltd. "
        f"| Confidential | {generated_date}"
    )
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in fp.runs:
        run.font.name      = "Arial"
        run.font.size      = Pt(9)
        run.font.color.rgb = RGB(0x00, 0x0B, 0x36)

    # ── Cover page ──────────────────────────────────────────────────────────
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_shading(cell, "000B36")
    tbl.rows[0].height = Inches(9)

    def cp(text, size, bold=False, color=None):
        p   = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name  = "Arial"
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.color.rgb = RGB(*color) if color else RGB(0xFF, 0xFF, 0xFF)

    cell.paragraphs[0].clear()
    cp("\n\nCOVVALENT", 24, bold=True)
    cp("ELN Intelligence — R&D Meeting Copilot", 18, color=(0x9D, 0xD1, 0xF1))
    cp(f"\n{topic}", 14, bold=True)
    if project_code:
        cp(f"Project: {project_code}", 12, color=(0x9D, 0xD1, 0xF1))
    cp(f"\n{generated_date} | {author or 'unknown'}", 11)
    cp("\nCONFIDENTIAL", 10, bold=True)
    doc.add_page_break()

    BODY_SECTIONS = [
        ("Context",           gpt_sections.get("context")),
        ("Hypothesis Review", gpt_sections.get("hypotheses")),
        ("Experiment Plan",   gpt_sections.get("experiment_plan")),
        ("Risks",             gpt_sections.get("risks")),
        ("Analysis",          gpt_sections.get("full")),
    ]

    for idx, (title, content) in enumerate(BODY_SECTIONS):
        if not content:
            continue
        h = doc.add_heading(f"Section {idx + 1}: {title}", level=1)
        for run in h.runs:
            run.font.name      = "Arial"
            run.font.color.rgb = RGB(0x00, 0x0B, 0x36)
        for para in content.split("\n\n"):
            if para.strip():
                bp = doc.add_paragraph(para.strip())
                for run in bp.runs:
                    run.font.name      = "Arial"
                    run.font.size      = Pt(11)
                    run.font.color.rgb = RGB(0x0E, 0x26, 0x73)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── Blob upload ───────────────────────────────────────────────────────────────
async def _upload_blob(content: bytes, blob_path: str, container: str = BLOB_CONTAINER) -> str:
    """
    Auto-creates `container` if missing, using the app's managed identity
    (already granted Storage Blob Data Contributor at the STORAGE ACCOUNT
    scope — this pattern extends to any container name, including the new
    eln-meeting-audio, with no new role assignment needed). Shared by
    eln-reports (docx reports + JSON sidecars) and eln-meeting-audio (raw
    chunk / whole-file audio archival).
    """
    from azure.identity.aio import DefaultAzureCredential
    from azure.storage.blob.aio import BlobServiceClient
    account_url = f"https://{STORAGE_ACCOUNT}.blob.core.windows.net"
    cred        = DefaultAzureCredential()
    try:
        async with BlobServiceClient(account_url=account_url, credential=cred) as svc:
            try:
                await svc.create_container(container)
            except Exception:
                pass
            await svc.get_blob_client(
                container=container, blob=blob_path
            ).upload_blob(content, overwrite=True)
    finally:
        await cred.close()
    return f"{account_url}/{container}/{blob_path}"


async def _download_blob_json(blob_path: str, container: str = BLOB_CONTAINER) -> Optional[dict]:
    """Read a small JSON sidecar blob's content directly (not a SAS
    redirect — this is for the backend to read its own sidecar, not for a
    browser to download it)."""
    from azure.identity.aio import DefaultAzureCredential
    from azure.storage.blob.aio import BlobServiceClient
    account_url = f"https://{STORAGE_ACCOUNT}.blob.core.windows.net"
    cred        = DefaultAzureCredential()
    try:
        async with BlobServiceClient(account_url=account_url, credential=cred) as svc:
            blob_client = svc.get_blob_client(container=container, blob=blob_path)
            stream      = await blob_client.download_blob()
            data        = await stream.readall()
            return json.loads(data)
    except Exception as e:
        logger.warning(f"[MEETING] Could not read sidecar blob {container}/{blob_path}: {e}")
        return None
    finally:
        await cred.close()

# ── Report-record helpers (status column only — no migration) ────────────────
def _insert_processing_report(project_code: Optional[str], author: str) -> int:
    """Inserts a placeholder row with status='processing' so
    GET /api/ai/meeting/status/{report_id} has something to poll
    immediately. Writes only to columns the original synchronous INSERT
    already used — topic/blob_url/transcript get real values once the
    background task finishes; placeholders here are just non-null."""
    conn = _get_conn()
    cur  = conn.cursor()
    cur.execute(
        "INSERT INTO eln_meeting_reports "
        "(project_code, topic, blob_url, transcript, author, status) "
        "OUTPUT INSERTED.report_id "
        "VALUES (?,?,?,?,?,'processing')",
        project_code, "Processing…", "", "", author
    )
    report_id = cur.fetchone()[0]
    conn.commit()
    conn.close()
    return report_id


def _update_meeting_report(report_id: int, status: str, topic: str = None, blob_url: str = None, transcript: str = None):
    """Only ever writes to columns the original code already wrote to
    (status, topic, blob_url, transcript) — no new columns, no migration."""
    conn = _get_conn()
    cur  = conn.cursor()
    if topic is not None and blob_url is not None and transcript is not None:
        cur.execute(
            "UPDATE eln_meeting_reports SET status=?, topic=?, blob_url=?, transcript=? WHERE report_id=?",
            status, topic, blob_url, transcript, report_id
        )
    elif topic is not None:
        cur.execute(
            "UPDATE eln_meeting_reports SET status=?, topic=? WHERE report_id=?",
            status, topic, report_id
        )
    else:
        cur.execute("UPDATE eln_meeting_reports SET status=? WHERE report_id=?", status, report_id)
    conn.commit()
    conn.close()

# ── Meeting copilot system prompt ─────────────────────────────────────────────
MEETING_SYSTEM = """
You are the R&D Meeting Co-Pilot for Covvalent, a specialty chemicals CDMO in India.
You receive R&D meeting transcripts in English, Hinglish (Hindi-English mix), or any
combination. Your output is ALWAYS in English regardless of input language.

Analyse the transcript and produce a structured report with EXACTLY these six
sections (use these exact headers):

## CONTEXT
One paragraph only. What process/step is being discussed, what problem exists,
what the team is trying to achieve. No hypotheses or analysis here.

## HYPOTHESIS REVIEW
Bulleted list. For each hypothesis in the transcript:
- Sound: [hypothesis statement] (Verdict: Sound) — no reasoning needed
- Off: [hypothesis] — 1-3 lines on what is wrong and why
- Needs sharpening: [hypothesis] — 1-3 lines reframing it precisely
Then add: Added by reviewer: [hypothesis the team missed] — 1-3 lines on
mechanism and plausibility

## EXPERIMENT PLAN
First: one paragraph stating DoE family choice and expected stop point.
Then numbered bullets: [Factor varied] | [Range/levels] | [Response measured] |
[Stop/proceed rule]

## RISKS
2-5 bullets only. Major safety risks of the PROPOSED experiments: exotherm,
pressure, gas evolution, reactive intermediates, runaway risk. One sentence each.

## CANDIDATE_NOTES
A JSON array (in a fenced ```json block) of project-memory candidates found in
the transcript. Two categories only:
- "decision": strategic/process decisions — route changes, abandoned
  approaches, procedural changes ("we decided to...", "dropping the...",
  "holding at...")
- "data_point": a reported measured value or correction not certain to be in
  the ELN yet — especially if described as approximate, recalled from memory,
  or contradicting a logged value.
Each item: {"note_type": "decision"|"data_point", "note_text": "...",
"exp_number_full": "..."|null, "project_code": "..."|null}
note_text must be the FULL, complete statement with real context — do not
truncate to a fake one-liner. Do not invent items not grounded in the
transcript. If nothing qualifies, return an empty array.

Resolving project codes: A PROJECT DIRECTORY appears before the transcript,
listing every project_code alongside its product name, generic name, and CAS
number. For every candidate note, resolve project_code by matching whatever
product is actually being discussed in that part of the transcript against the
directory — including informal references ("the sulfonation project",
"tryptophan work"). Different candidates in the same meeting may belong to
different projects if more than one product was discussed — resolve each one
independently, do not apply one project_code to all candidates by default. If
no directory entry matches with reasonable confidence, set project_code to
null. Do not guess a project_code from a vague or ambiguous reference.

## EMAIL_DRAFT
A JSON object (in a fenced ```json block): {"subject": "...", "body": "..."}
- subject: "Meeting Summary — <topic/project(s)> — <date>"
  Use the PROJECT DIRECTORY the same way to name every project actually
  discussed in the subject line (e.g. "Meeting Summary — O031C00 / H000E02 —
  <date>"), not just the project_code the request was filed under, since a
  meeting may cover more than one product.
- body:
  1. One short paragraph (1-3 sentences) framing what was discussed.
  2. "Decisions & Next Steps:" — numbered list. Each item states the
     decision/finding in one sentence, then "-> " followed by the specific
     actionable next step.
  3. One closing sentence on safety flags, or state none were raised.
  4. Keep the body under ~280 words total (roughly a 2-minute read). If there
     are more than 5 decisions, include only the 5 most consequential and add
     "(N additional items covered in the full report)".

Language rules: Read Hinglish naturally. Filler words (bhai, yaar, na, toh,
achha) carry no chemistry — ignore them. Technical terms in English carry the
meaning. Output in English.
"""

# ── Shared report-generation pipeline (GPT -> sections -> docx -> blob) ──────
async def _generate_meeting_report_from_transcript(
    transcript:   str,
    project_code: Optional[str],
    author:       str,
) -> dict:
    """
    Shared by the synchronous /api/ai/meeting endpoint and the
    /api/ai/meeting/upload background task. Blob paths are namespaced by a
    slugified author so concurrent meetings from different users never
    collide. Writes a JSON sidecar (candidate_notes, email_draft) next to
    the report blob — same container, same key prefix, .json suffix — since
    no new SQL columns/migration are allowed for these fields.
    """
    conn = _get_conn()
    directory = get_project_directory_cached(conn)
    conn.close()
    logger.info(f"Loaded {len(directory)} projects for resolution")

    directory_json = json.dumps(directory, ensure_ascii=False)
    user_content = (
        "PROJECT DIRECTORY (for resolving product names mentioned in the transcript to "
        "their project_code — match on product name, generic name, or CAS number, even "
        "if referenced informally):\n"
        f"{directory_json}\n\n"
        f"TRANSCRIPT:\n\n{transcript[:8000]}"
    )

    gpt_output = await _call_gpt(MEETING_SYSTEM, user_content)

    sections       = _parse_sections(gpt_output)
    topic          = _extract_topic(gpt_output)
    generated_date = datetime.utcnow().strftime("%Y-%m-%d")
    slug           = _slug(topic)
    author_slug    = _slug(author or "unknown", max_len=30) or "unknown"
    filename       = f"{generated_date}-{slug}-meeting.docx"
    blob_path      = f"meetings/{generated_date}/{author_slug}/{filename}"

    candidate_notes, email_draft = _parse_candidate_notes_and_email(sections)

    docx_bytes = _build_meeting_report_docx(topic, project_code, author, sections, generated_date)
    blob_url   = await _upload_blob(docx_bytes, blob_path)

    sidecar_path = blob_path[:-len(".docx")] + ".json"
    try:
        await _upload_blob(
            json.dumps(
                {"candidate_notes": candidate_notes, "email_draft": email_draft, "author": author},
                ensure_ascii=False, indent=2
            ).encode("utf-8"),
            sidecar_path,
        )
    except Exception as e:
        logger.warning(f"[MEETING] Failed to write sidecar JSON at {sidecar_path}: {e}")

    return {
        "topic":            topic,
        "generated_date":   generated_date,
        "filename":         filename,
        "blob_url":         blob_url,
        "transcript":       transcript,
        "candidate_notes":  candidate_notes,
        "email_draft":      email_draft,
    }


async def _process_meeting_upload(report_id: int, audio_bytes: bytes, ext: str, content_type: str, project_code: Optional[str], author: str):
    """
    Background task for POST /api/ai/meeting/upload — runs AFTER the HTTP
    response has already been sent. A multi-hour transcription + report
    pipeline here never risks the App Service front-end idle timeout that
    broke meetings past ~1.5 hours under the old synchronous SDK-based
    flow, because nothing here is holding an HTTP response open.
    """
    author_slug = _slug(author or "unknown", max_len=30) or "unknown"

    # Archive the raw uploaded audio, namespaced by author + report_id so
    # concurrent uploads from different users never collide. Non-fatal if
    # it fails — still attempt transcription either way.
    try:
        await _upload_blob(audio_bytes, f"{author_slug}/{report_id}/original{ext}", container=AUDIO_CONTAINER)
    except Exception as e:
        logger.warning(f"[MEETING UPLOAD] Failed to archive original audio for report {report_id}: {e}")

    try:
        result = await _transcribe_via_rest(audio_bytes, f"upload{ext}", content_type=content_type, timeout=600.0)
        full_text, phrase_locales = _extract_transcript_and_locales(result)

        if not full_text.strip():
            _update_meeting_report(report_id, status="failed", topic="Transcription produced no text", blob_url="", transcript="")
            return

        report = await _generate_meeting_report_from_transcript(full_text, project_code, author)

        for item in report["candidate_notes"]:
            if not item.get("project_code"):
                item["project_code"] = project_code
            item["source_report_id"] = report_id

        # Fold the per-chunk-equivalent phrase-locale detail into the same
        # sidecar the shared pipeline just wrote, rather than a second blob.
        sidecar_path = f"meetings/{report['generated_date']}/{author_slug}/" \
                       f"{report['filename'][:-len('.docx')]}.json"
        try:
            await _upload_blob(
                json.dumps(
                    {
                        "candidate_notes": report["candidate_notes"],
                        "email_draft":     report["email_draft"],
                        "author":          author,
                        "phrase_locales":  phrase_locales,
                    },
                    ensure_ascii=False, indent=2
                ).encode("utf-8"),
                sidecar_path,
            )
        except Exception as e:
            logger.warning(f"[MEETING UPLOAD] Failed to write sidecar JSON for report {report_id}: {e}")

        _update_meeting_report(
            report_id, status="complete",
            topic=report["topic"][:200], blob_url=report["blob_url"],
            transcript=full_text[:8000],
        )
    except Exception:
        logger.exception(f"[MEETING UPLOAD] Background processing failed for report {report_id}")
        try:
            _update_meeting_report(report_id, status="failed")
        except Exception:
            logger.exception(f"[MEETING UPLOAD] Failed to mark report {report_id} as failed")

# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/api/ai/speech")
async def transcribe_speech(
    audio:      Optional[UploadFile] = File(None),
    transcript: Optional[str]        = Form(None),
):
    """Transcribe audio via the Fast Transcription REST API, or pass through a text transcript."""
    if transcript:
        return {"transcript": transcript, "language_detected": "text"}
    if not audio:
        raise HTTPException(400, detail="Provide either an audio file or a transcript field.")

    audio_bytes  = await audio.read()
    content_type = audio.content_type or "audio/wav"

    try:
        text, lang = await _transcribe_audio(audio_bytes, content_type)
    except httpx.HTTPStatusError as e:
        raise HTTPException(502, detail=f"Transcription failed: {e}")
    except Exception as e:
        raise HTTPException(500, detail=f"Transcription failed: {e}")

    return {"transcript": text, "language_detected": lang}


@router.post("/api/ai/meeting/chunk", response_model=ChunkResponse)
async def transcribe_meeting_chunk(
    audio:       UploadFile = File(...),
    session_id:  str        = Form(...),
    chunk_index: int        = Form(...),
    author:      str        = Form(""),
):
    """
    Transcribes ONE short (~4 min) segment of a live in-progress recording.
    A single short-audio Fast Transcription call is fast and well within any
    App Service front-end timeout — this is what lets a multi-hour live
    meeting survive without one giant inline call. Also durably archives the
    raw chunk audio to eln-meeting-audio, namespaced by author + session_id,
    so two concurrent meetings from different people never collide.
    """
    audio_bytes = await audio.read()
    if not audio_bytes:
        raise HTTPException(400, detail="Empty audio chunk.")

    content_type = audio.content_type or "audio/webm"
    ext = _EXT_MAP.get(content_type.split(";")[0].strip().lower(), ".webm")

    author_slug = _slug(author or "unknown", max_len=30) or "unknown"
    blob_path   = f"{author_slug}/{session_id}/chunk-{chunk_index:04d}{ext}"

    try:
        await _upload_blob(audio_bytes, blob_path, container=AUDIO_CONTAINER)
    except Exception as e:
        # Non-fatal — still attempt transcription even if archival failed.
        logger.warning(f"[MEETING CHUNK] Failed to archive chunk audio ({blob_path}): {e}")

    try:
        result = await _transcribe_via_rest(audio_bytes, f"chunk{ext}", content_type=content_type)
    except httpx.HTTPStatusError as e:
        raise HTTPException(502, detail=f"Transcription failed for chunk {chunk_index}: {e}")
    except Exception as e:
        raise HTTPException(500, detail=f"Transcription failed for chunk {chunk_index}: {e}")

    full_text, phrase_locales = _extract_transcript_and_locales(result)
    locale = next((p["locale"] for p in phrase_locales if p.get("locale")), None)

    return ChunkResponse(session_id=session_id, chunk_index=chunk_index, text=full_text, locale=locale)


@router.post("/api/ai/meeting/upload", response_model=MeetingUploadResponse, status_code=202)
async def upload_meeting_recording(
    background_tasks: BackgroundTasks,
    audio:            UploadFile    = File(...),
    project_code:     Optional[str] = Form(None),
    author:           str           = Form(""),
):
    """
    Accepts one whole pre-recorded meeting file. Inserts a 'processing' row
    immediately and returns — transcription (which can take minutes for a
    long file) and report generation both run in a BackgroundTask, AFTER
    the HTTP response is sent, so this never risks the App Service
    front-end idle timeout the old synchronous transcribe-then-generate
    flow hit past ~1.5 hours of audio.
    """
    audio_bytes = await audio.read()
    if not audio_bytes:
        raise HTTPException(400, detail="Uploaded audio file is empty.")

    content_type = audio.content_type or "audio/webm"
    ext = _EXT_MAP.get(content_type.split(";")[0].strip().lower(), ".webm")

    report_id = _insert_processing_report(project_code, author)

    background_tasks.add_task(
        _process_meeting_upload, report_id, audio_bytes, ext, content_type, project_code, author
    )

    return MeetingUploadResponse(report_id=report_id, status="processing")


@router.get("/api/ai/meeting/status/{report_id}")
async def get_meeting_status(report_id: int):
    """
    Poll target for POST /api/ai/meeting/upload. Once status is no longer
    'processing', also returns the sidecar's candidate_notes/email_draft —
    read from the JSON file next to the report blob (never stored in SQL,
    per the no-migration constraint).
    """
    conn = _get_conn()
    cur  = conn.cursor()
    cur.execute(
        "SELECT report_id, project_code, topic, blob_url, author, generated_date, status "
        "FROM eln_meeting_reports WHERE report_id = ?",
        report_id
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, detail=f"No meeting report with report_id={report_id}")

    result = {
        "report_id":       row[0],
        "project_code":    row[1],
        "topic":           row[2],
        "blob_url":        row[3],
        "author":          row[4],
        "generated_date":  _serialize(row[5]),
        "status":          row[6],
        "candidate_notes": [],
        "email_draft":     {"subject": "", "body": ""},
    }

    if result["status"] == "complete" and result["blob_url"]:
        match = re.search(rf"{re.escape(BLOB_CONTAINER)}/(.+)\.docx$", result["blob_url"])
        if match:
            sidecar_path = match.group(1) + ".json"
            sidecar = await _download_blob_json(sidecar_path)
            if sidecar:
                result["candidate_notes"] = sidecar.get("candidate_notes", [])
                result["email_draft"]     = sidecar.get("email_draft", {"subject": "", "body": ""})

    return result


@router.post("/api/ai/meeting")
async def generate_meeting_report(req: MeetingRequest):
    """Run meeting transcript through copilot prompt, generate Word report, upload to blob.
    External contract unchanged."""
    if not req.transcript.strip():
        raise HTTPException(400, detail="Transcript is empty.")

    report = await _generate_meeting_report_from_transcript(req.transcript, req.project_code, req.author)

    conn = _get_conn()
    cur  = conn.cursor()
    cur.execute(
        "INSERT INTO eln_meeting_reports "
        "(project_code, topic, blob_url, transcript, author, status) "
        "OUTPUT INSERTED.report_id, INSERTED.generated_date "
        "VALUES (?,?,?,?,?,'complete')",
        req.project_code, report["topic"][:200], report["blob_url"], req.transcript[:8000], req.author
    )
    row = cur.fetchone()
    report_id  = row[0]
    db_date    = _serialize(row[1])
    conn.commit()
    conn.close()

    for item in report["candidate_notes"]:
        # Priority: (1) model's directory-resolved project_code, (2) the
        # meeting's own project_code param if the model returned null,
        # (3) otherwise stays null.
        if not item.get("project_code"):
            item["project_code"] = req.project_code
        item["source_report_id"] = report_id

    return {
        "report_id":       report_id,
        "blob_url":        report["blob_url"],
        "topic":           report["topic"],
        "generated_date":  db_date,
        "filename":        report["filename"],
        "candidate_notes": report["candidate_notes"],
        "email_draft":     report["email_draft"],
    }


@router.get("/api/ai/meeting/reports")
def get_meeting_reports():
    """All meeting reports, newest first."""
    conn = _get_conn()
    cur  = conn.cursor()
    cur.execute(
        "SELECT report_id, project_code, topic, blob_url, author, generated_date, status "
        "FROM eln_meeting_reports ORDER BY generated_date DESC"
    )
    reports = _clean(_rows_to_dicts(cur))
    conn.close()
    return {"reports": reports}
