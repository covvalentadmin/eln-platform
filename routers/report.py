"""
routers/report.py — ELN Project Report Generation + Project Notes
POST /api/ai/report        — async project analysis report (202 Accepted)
GET  /api/ai/report/{code} — list reports for a project
GET  /api/ai/reports       — all reports across all projects
POST /api/ai/notes         — create project note
GET  /api/ai/notes/{code}  — list active notes for a project
"""

import os
import io
import json
import asyncio
import traceback
import httpx
import pyodbc
from datetime import datetime, date
from typing import Optional
from fastapi import APIRouter, BackgroundTasks, HTTPException
from pydantic import BaseModel

router = APIRouter()

# ── Config ────────────────────────────────────────────────────────────────────
AOAI_ENDPOINT    = os.environ.get("AOAI_ENDPOINT", "https://aoai-eln-covvalent-2e2ec.openai.azure.com")
AOAI_DEPLOYMENT  = "gpt-4o"
AOAI_API_VERSION = "2024-12-01-preview"
STORAGE_ACCOUNT  = "stelncoovalent"
BLOB_CONTAINER   = "eln-reports"

API_BASE = os.environ.get(
    "INTERNAL_API_BASE",
    "https://eln-api-covvalent-asfhf0abbvh2bphd.southindia-01.azurewebsites.net"
)

# ── Pydantic models ───────────────────────────────────────────────────────────
class ReportRequest(BaseModel):
    project_code: str
    triggered_by: str = "manual"

class NoteRequest(BaseModel):
    project_code: str
    note_text: str
    author: str
    captured_from: str = "manual"

def _extract_summary(analysis_text: str, max_chars: int = 800) -> str:
    """
    Extract executive summary paragraph from GPT-4o analysis output.
    Stored in SQL for the agent tool — must be concise.
    """
    import re
    patterns = [
        r"(?:executive summary|overview|summary)[:\s]*\n+([\s\S]+?)(?:\n\n|\n#+|\Z)",
        r"#+\s*(?:executive summary|overview|summary)\s*\n+([\s\S]+?)(?:\n\n|\n#+|\Z)",
    ]
    for pat in patterns:
        m = re.search(pat, analysis_text, re.IGNORECASE)
        if m:
            candidate = m.group(1).strip()
            if len(candidate) > 80:
                return candidate[:max_chars]
    paragraphs = [p.strip() for p in analysis_text.split("\n\n") if p.strip()]
    for para in paragraphs:
        if len(para) > 100:
            return para[:max_chars]
    return analysis_text[:max_chars]

# ── DB helpers ────────────────────────────────────────────────────────────────
def _get_conn():
    conn_str = (
        f"DRIVER={{ODBC Driver 18 for SQL Server}};"
        f"SERVER={os.environ['SQL_SERVER']};"
        f"DATABASE={os.environ['SQL_DATABASE']};"
        f"UID={os.environ['SQL_USERNAME']};"
        f"PWD={os.environ['SQL_PASSWORD']};"
        f"TrustServerCertificate=yes;"
        f"Connection Timeout=30;"
    )
    return pyodbc.connect(conn_str)

def _serialize(v):
    if isinstance(v, (datetime, date)):
        return v.isoformat()
    return v

def _clean(rows):
    return [{k: _serialize(v) for k, v in row.items()} for row in rows]

def _rows_to_dicts(cur):
    cols = [c[0] for c in cur.description]
    return [dict(zip(cols, row)) for row in cur.fetchall()]

# ── Auth helper ───────────────────────────────────────────────────────────────
async def _get_aoai_token() -> str:
    from azure.identity.aio import DefaultAzureCredential
    cred = DefaultAzureCredential()
    token = await cred.get_token("https://cognitiveservices.azure.com/.default")
    await cred.close()
    return token.token

# ── GPT call ──────────────────────────────────────────────────────────────────
async def _call_gpt(system_prompt: str, user_content: str) -> str:
    token = await _get_aoai_token()
    async with httpx.AsyncClient() as client:
        r = await client.post(
            f"{AOAI_ENDPOINT}/openai/deployments/{AOAI_DEPLOYMENT}/chat/completions?api-version={AOAI_API_VERSION}",
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            json={
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": user_content},
                ],
                "max_tokens": 4096,
                "temperature": 0.3,
            },
            timeout=120.0,
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]

# ── Word doc helpers ──────────────────────────────────────────────────────────
def _set_cell_shading(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def _add_heading(doc, text: str, level: int = 1):
    from docx.shared import RGBColor
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x00, 0x0B, 0x36)

def _add_body(doc, text: str):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x0E, 0x26, 0x73)

def _build_project_report_docx(
    project: dict,
    experiments: list,
    gpt_analysis: str,
    literature: dict,
    generated_date: str,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # Footer
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = f"Covvalent / Rainboweucalyptus Technologies Pvt. Ltd. | Confidential | {generated_date}"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x4A, 0x61, 0x94)

    # ── Cover page ──────────────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "000B36")
    table.rows[0].height = Inches(9)

    def _cover_para(cell, text, size, bold=False, color=None):
        from docx.shared import RGBColor as RGB
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGB(*color) if color else RGB(0xFF, 0xFF, 0xFF)

    cell.paragraphs[0].clear()
    _cover_para(cell, "\n\nCOVVALENT", 32, bold=True)
    _cover_para(cell, "ELN Intelligence — Project Analysis Report", 16, color=(0x9D, 0xD1, 0xF1))
    _cover_para(cell, f"\n{project.get('project_code','')} — {project.get('project_title','')}", 14, bold=True)
    if project.get("cas_number"):
        _cover_para(cell, f"CAS {project['cas_number']}", 12, color=(0x9D, 0xD1, 0xF1))
    _cover_para(cell, f"\n{generated_date}", 11)
    _cover_para(cell, "\nCONFIDENTIAL", 10, bold=True, color=(0xFF, 0xFF, 0xFF))

    doc.add_page_break()

    # ── Section 1: Project Identity ──────────────────────────────────────────
    _add_heading(doc, "Section 1: Project Identity")
    fields = [
        ("Product Name",    project.get("project_title", "—")),
        ("CAS Number",      project.get("cas_number", "—")),
        ("Generic Name",    project.get("generic_name", "—")),
        ("IUPAC Name",      project.get("iupac_name", "—")),
        ("Status",          "Active" if project.get("project_status") == 1 else "Closed"),
        ("Start Date",      (project.get("start_date") or "—")[:10]),
        ("Total Experiments", str(len(experiments))),
    ]
    from docx.shared import RGBColor as RGB
    for label, value in fields:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.color.rgb = RGB(0x00, 0x0B, 0x36)
        r2 = p.add_run(value)
        r2.font.name = "Arial"
        r2.font.color.rgb = RGB(0x0E, 0x26, 0x73)

    doc.add_page_break()

    # ── Section 2: Experiment History ────────────────────────────────────────
    _add_heading(doc, "Section 2: Experiment History")

    if experiments:
        headers = ["Exp Number", "Date", "Author", "Title", "Status"]
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            _set_cell_shading(hdr[i], "000B36")
            run = hdr[i].paragraphs[0].runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGB(0xFF, 0xFF, 0xFF)

        for idx, exp in enumerate(experiments):
            row = tbl.add_row()
            values = [
                exp.get("exp_number_full", ""),
                (exp.get("created_date") or "")[:10],
                exp.get("author") or "",
                exp.get("title") or "",
                str(exp.get("experiment_status", "")),
            ]
            bg = "FFFFFF" if idx % 2 == 0 else "DEEBF7"
            for j, val in enumerate(values):
                row.cells[j].text = val
                _set_cell_shading(row.cells[j], bg)
                row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)
                row.cells[j].paragraphs[0].runs[0].font.name = "Arial"
    else:
        _add_body(doc, "No experiments found for this project.")

    doc.add_page_break()

    # ── Section 3: Chemistry Analysis (GPT) ──────────────────────────────────
    _add_heading(doc, "Section 3: Chemistry Analysis")
    _add_heading(doc, "AI-Generated Analysis (gpt-4o)", level=2)
    for paragraph in gpt_analysis.split("\n\n"):
        if paragraph.strip():
            _add_body(doc, paragraph.strip())

    doc.add_page_break()

    # ── Section 4: Literature & Patent Context ────────────────────────────────
    _add_heading(doc, "Section 4: Literature & Patent Context")

    pubchem_raw = literature.get("pubchem") or {}
    pubchem = pubchem_raw[0] if isinstance(pubchem_raw, list) and pubchem_raw else pubchem_raw if isinstance(pubchem_raw, dict) else {}
    if pubchem:
        _add_heading(doc, "PubChem Compound Profile", level=2)
        for key in ["IUPACName", "MolecularFormula", "MolecularWeight", "InChIKey"]:
            val = pubchem.get(key)
            if val:
                p = doc.add_paragraph()
                r1 = p.add_run(f"{key}: ")
                r1.bold = True; r1.font.name = "Arial"; r1.font.color.rgb = RGB(0x00, 0x0B, 0x36)
                r2 = p.add_run(str(val))
                r2.font.name = "Arial"; r2.font.color.rgb = RGB(0x0E, 0x26, 0x73)

    papers = literature.get("papers") or []
    if papers:
        _add_heading(doc, "CrossRef Papers", level=2)
        for paper in papers[:10]:
            title = paper.get("title", [""])[0] if isinstance(paper.get("title"), list) else paper.get("title", "")
            doi   = paper.get("DOI", "")
            year  = paper.get("published-print", {}).get("date-parts", [[""]])[0][0] if paper.get("published-print") else ""
            p = doc.add_paragraph(f"• {title} ({year})")
            if doi:
                p.add_run(f" — DOI: {doi}")
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.color.rgb = RGB(0x0E, 0x26, 0x73)

    if not pubchem and not papers:
        _add_body(doc, "No literature data retrieved for this project.")

    doc.add_page_break()

    # ── Section 5: Recommendations ────────────────────────────────────────────
    _add_heading(doc, "Section 5: Recommendations")
    _add_body(doc, "Priority-ordered next experiments are included in Section 3 under RECOMMENDATIONS. "
                   "Refer to the AI analysis above for explicit stop rules and experimental queue.")

    # Save
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── Blob upload ───────────────────────────────────────────────────────────────
async def _upload_blob(content: bytes, blob_path: str) -> str:
    from azure.identity.aio import DefaultAzureCredential
    from azure.storage.blob.aio import BlobServiceClient
    account_url = f"https://{STORAGE_ACCOUNT}.blob.core.windows.net"
    credential = DefaultAzureCredential()
    try:
        async with BlobServiceClient(account_url=account_url, credential=credential) as svc:
            try:
                await svc.create_container(BLOB_CONTAINER)
            except Exception:
                pass
            blob = svc.get_blob_client(container=BLOB_CONTAINER, blob=blob_path)
            await blob.upload_blob(content, overwrite=True)
    finally:
        await credential.close()
    return f"{account_url}/{BLOB_CONTAINER}/{blob_path}"

# ── Report background task ────────────────────────────────────────────────────
def _update_report_record(report_id: int, status: str, blob_url: str = None, exp_count: int = None, report_summary: str = None):
    conn = _get_conn()
    cur  = conn.cursor()
    if blob_url and exp_count is not None:
        cur.execute(
            "UPDATE eln_project_reports SET status=?, blob_url=?, experiment_count=?, report_summary=? WHERE report_id=?",
            status, blob_url, exp_count, report_summary, report_id
        )
    else:
        cur.execute("UPDATE eln_project_reports SET status=? WHERE report_id=?", status, report_id)
    conn.commit()
    conn.close()

async def _run_report_generation(
    report_id: int,
    project_code: str,
    generated_by: str,
):
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            # Fetch project + experiments
            r = await client.post(f"{API_BASE}/api/ai/fetch", json={"project_code": project_code})
            r.raise_for_status()
            fetch_data = r.json()

        project_info     = fetch_data.get("project", {})
        experiments      = fetch_data.get("experiments", [])
        product_name     = project_info.get("project_title", project_code)
        cas_number       = project_info.get("cas_number", "unknown")
        experiment_count = len(experiments)

        # Fetch literature
        literature = {}
        if cas_number and cas_number != "unknown":
            try:
                async with httpx.AsyncClient(timeout=30.0) as client:
                    r = await client.get(f"{API_BASE}/api/ai/literature", params={"q": cas_number})
                    if r.status_code == 200:
                        literature = r.json()
            except Exception:
                pass

        # Chemistry analysis prompt
        CHEMISTRY_SYSTEM = (
            "You are a senior process chemist reviewing Covvalent's internal experiment history. "
            "Use lab-report language. Cite exp_number_full for every factual claim. Never fabricate data."
        )
        exp_summary = "\n".join(
            f"- {e.get('exp_number_full')} ({(e.get('created_date') or '')[:10]}): {e.get('title','')} "
            f"[{e.get('author','')}] status={e.get('experiment_status','')}"
            for e in experiments[:30]
        )
        CHEMISTRY_USER = f"""You are reviewing Covvalent's internal experiment history for {product_name} (CAS {cas_number}).

Experiment data:
{exp_summary}

Generate a structured analysis with these FIVE numbered sections:

1. REACTION MECHANISM: Identify the reaction class and explain the molecular-level pathway. Why do the key conditions (temperature, pH, reagents) matter mechanistically?
2. CRITICAL PROCESS PARAMETERS: List each CPP with its safe operating range and consequence of violation.
3. YIELD & PURITY TRENDS: What patterns emerge across experiments? What drove the best results?
4. GAPS: What analytical data is missing? What conditions have not been explored?
5. RECOMMENDATIONS: Priority-ordered queue of next experiments with explicit stop rules. Analytical characterisation first, then variable isolation, then optimisation."""

        gpt_analysis = await _call_gpt(CHEMISTRY_SYSTEM, CHEMISTRY_USER)

        # Build Word doc
        generated_date = datetime.utcnow().strftime("%Y-%m-%d")
        docx_bytes = _build_project_report_docx(
            project_info, experiments, gpt_analysis, literature, generated_date
        )

        # Upload to blob
        blob_path = f"{project_code}/{generated_date}-{project_code}-analysis.docx"
        blob_url  = await _upload_blob(docx_bytes, blob_path)

        report_summary = _extract_summary(gpt_analysis)
        _update_report_record(report_id, "complete", blob_url, experiment_count, report_summary)

    except Exception as e:
        error_msg = traceback.format_exc()
        print(f"[REPORT BACKGROUND ERROR] report_id={report_id}: {error_msg}")
        try:
            conn = _get_conn()
            cur  = conn.cursor()
            cur.execute(
                "UPDATE eln_project_reports SET status='failed', blob_url=? WHERE report_id=?",
                f"ERROR: {str(e)[:400]}", report_id
            )
            conn.commit()
            conn.close()
        except Exception as db_err:
            print(f"[REPORT BACKGROUND ERROR] DB update also failed: {db_err}")

# ── Endpoints ─────────────────────────────────────────────────────────────────


# ── Download endpoint ─────────────────────────────────────────────────────────
import re as _re
import datetime as _dt
from azure.storage.blob import generate_blob_sas, BlobSasPermissions
from azure.storage.blob import BlobServiceClient as _BlobServiceClient
from azure.identity import ManagedIdentityCredential as _MIC
from fastapi.responses import RedirectResponse

_STORAGE_ACCOUNT  = "stelncoovalent"
_REPORTS_CONTAINER = "eln-reports"

@router.get("/api/ai/report/download/{report_id}")
def download_report(report_id: int):
    """Generate a 1-hour SAS URL for a report blob and redirect."""
    conn = _get_conn()
    try:
        row = conn.execute(
            "SELECT blob_url FROM eln_project_reports WHERE report_id=?",
            (report_id,)
        ).fetchone()
    finally:
        conn.close()

    if not row or not row[0] or not row[0].startswith("https://"):
        raise HTTPException(status_code=404, detail="Report not found or not ready")

    blob_url = row[0]
    match = _re.search(rf"{_REPORTS_CONTAINER}/(.+)$", blob_url)
    if not match:
        raise HTTPException(status_code=500, detail="Cannot parse blob path from URL")
    blob_path = match.group(1)

    credential = _MIC()
    blob_service = _BlobServiceClient(
        account_url=f"https://{_STORAGE_ACCOUNT}.blob.core.windows.net",
        credential=credential
    )
    start  = _dt.datetime.utcnow() - _dt.timedelta(minutes=5)
    expiry = _dt.datetime.utcnow() + _dt.timedelta(hours=1)
    delegation_key = blob_service.get_user_delegation_key(start, expiry)

    sas = generate_blob_sas(
        account_name=_STORAGE_ACCOUNT,
        container_name=_REPORTS_CONTAINER,
        blob_name=blob_path,
        user_delegation_key=delegation_key,
        permission=BlobSasPermissions(read=True),
        expiry=expiry
    )
    sas_url = f"https://{_STORAGE_ACCOUNT}.blob.core.windows.net/{_REPORTS_CONTAINER}/{blob_path}?{sas}"
    return RedirectResponse(url=sas_url)

@router.post("/api/ai/report", status_code=202)
async def generate_report(req: ReportRequest, background_tasks: BackgroundTasks):
    """Kick off project analysis report generation. Returns 202 immediately."""
    conn = _get_conn()
    cur  = conn.cursor()
    cur.execute(
        "INSERT INTO eln_project_reports (project_code, triggered_by, generated_by, status) "
        "OUTPUT INSERTED.report_id, INSERTED.generated_date VALUES (?,?,?,'pending')",
        req.project_code, req.triggered_by, req.triggered_by
    )
    row = cur.fetchone()
    report_id      = row[0]
    generated_date = _serialize(row[1])
    conn.commit()
    conn.close()

    background_tasks.add_task(_run_report_generation, report_id, req.project_code, req.triggered_by)

    return {
        "report_id":      report_id,
        "project_code":   req.project_code,
        "status":         "pending",
        "generated_date": generated_date,
        "message":        "Report generation started. Poll GET /api/ai/report/{project_code} for status.",
    }


@router.get("/api/ai/report/list/{project_code}")
def list_project_reports(project_code: str):
    """
    Agent tool endpoint — reports for one project, newest first.
    Lightweight: no blob reads, no SAS generation.
    """
    api_base = os.environ.get(
        "API_BASE",
        "https://eln-api-covvalent-asfhf0abbvh2bphd.southindia-01.azurewebsites.net",
    )
    try:
        conn = _get_conn()
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT report_id, project_code, generated_date, experiment_count,
                   status, report_summary
            FROM eln_project_reports
            WHERE project_code = ?
            ORDER BY generated_date DESC
            """,
            project_code,
        )
        rows = cursor.fetchall()
        conn.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    return {
        "project_code": project_code,
        "report_count": len(rows),
        "reports": [
            {
                "report_id":        r.report_id,
                "generated_date":   str(r.generated_date),
                "experiment_count": r.experiment_count,
                "status":           r.status,
                "report_summary":   r.report_summary,
                "download_url": (
                    f"{api_base}/api/ai/report/download/{r.report_id}"
                    if r.status == "complete" else None
                ),
            }
            for r in rows
        ],
    }


@router.get("/api/ai/report/{project_code}")
def get_project_reports(project_code: str):
    """List all reports for a project."""
    conn = _get_conn()
    cur  = conn.cursor()
    cur.execute(
        "SELECT report_id, generated_date, blob_url, status, experiment_count, triggered_by, generated_by "
        "FROM eln_project_reports WHERE project_code = ? ORDER BY generated_date DESC",
        project_code
    )
    reports = _clean(_rows_to_dicts(cur))
    conn.close()
    return {"project_code": project_code, "reports": reports}


@router.get("/api/ai/reports")
def get_all_reports():
    """All reports across all projects, newest first."""
    conn = _get_conn()
    cur  = conn.cursor()
    cur.execute(
        "SELECT report_id, project_code, generated_date, blob_url, status, "
        "experiment_count, triggered_by, generated_by "
        "FROM eln_project_reports ORDER BY generated_date DESC"
    )
    reports = _clean(_rows_to_dicts(cur))
    conn.close()
    return {"reports": reports}


@router.post("/api/ai/notes", status_code=201)
def create_note(req: NoteRequest):
    """Create a project note."""
    conn = _get_conn()
    cur  = conn.cursor()
    cur.execute(
        "INSERT INTO eln_project_notes (project_code, note_text, captured_from, author) "
        "OUTPUT INSERTED.note_id, INSERTED.created_date VALUES (?,?,?,?)",
        req.project_code, req.note_text, req.captured_from, req.author
    )
    row = cur.fetchone()
    note_id      = row[0]
    created_date = _serialize(row[1])
    conn.commit()
    conn.close()
    return {"note_id": note_id, "project_code": req.project_code, "created_date": created_date}


@router.get("/api/ai/notes/{project_code}")
def get_notes(project_code: str):
    """All active notes for a project."""
    conn = _get_conn()
    cur  = conn.cursor()
    cur.execute(
        "SELECT note_id, note_text, author, created_date, captured_from "
        "FROM eln_project_notes WHERE project_code = ? AND is_deleted = 0 "
        "ORDER BY created_date DESC",
        project_code
    )
    notes = _clean(_rows_to_dicts(cur))
    conn.close()
    return {"project_code": project_code, "notes": notes}
